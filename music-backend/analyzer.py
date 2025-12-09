import librosa
import numpy as np
import subprocess
import os
import time
import json
import mutagen
import google.generativeai as genai
from basic_pitch.inference import predict
from docx import Document
from docx.shared import Pt


NOTES = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']
def _build_chord_templates(include_sevenths=True):
    """
    Build simple normalized templates for triads (and optionally 7th chords).
    Returns: dict: chord_name -> np.array(shape=(12,))
    """
    templates = {}

    def tpl(root, intervals):
        v = np.zeros(12, dtype=float)
        for i in intervals:
            v[(root + i) % 12] = 1.0
        return v

    for root in range(12):
        root_name = NOTES[root]

        # Triads
        templates[f"{root_name}"]    = tpl(root, [0, 4, 7])   # major
        templates[f"{root_name}m"]   = tpl(root, [0, 3, 7])   # minor
        templates[f"{root_name}dim"] = tpl(root, [0, 3, 6])   # diminished
        templates[f"{root_name}aug"] = tpl(root, [0, 4, 8])   # augmented

        if include_sevenths:
            templates[f"{root_name}7"]    = tpl(root, [0, 4, 7, 10])  # dominant 7
            templates[f"{root_name}maj7"] = tpl(root, [0, 4, 7, 11])  # major 7
            templates[f"{root_name}m7"]   = tpl(root, [0, 3, 7, 10])  # minor 7

    # Normalize for cosine similarity
    for k, v in templates.items():
        n = np.linalg.norm(v)
        if n > 0:
            templates[k] = v / n

    return templates


CHORD_TEMPLATES = _build_chord_templates(include_sevenths=True)

def _enforce_min_duration(chords, min_chord_duration):
    """
    Merge very short chord segments into the previous one.
    This makes chords look more stable in time.
    """
    if not chords:
        return []

    merged = [chords[0]]

    for seg in chords[1:]:
        dur = seg["end_time"] - seg["start_time"]

        if dur < min_chord_duration:
            # Too short: merge into previous segment by extending its end
            merged[-1]["end_time"] = max(merged[-1]["end_time"], seg["end_time"])
        else:
            merged.append(seg)

    return merged

def analyze_chords(
    file_path,
    hop_length=2048,
    chord_threshold=0.2,
    min_chord_duration=0.5,
    use_hpss=True
):
    """
    Detect chords using Librosa chroma + template matching.

    Returns:
        List[dict]: [{ "start_time": float, "end_time": float, "chord_name": str }, ...]
        or [{ "error": str }] on failure.
    """
    try:
        if not os.path.exists(file_path):
            return [{"error": f"File not found: {file_path}"}]

        # 1. Load
        y, sr = librosa.load(file_path, mono=True)
        if y.size == 0:
            return [{"error": "Audio file appears to be empty."}]

        # 2. HPSS
        if use_hpss and y.size >= 4096:
            try:
                y_harm, _ = librosa.effects.hpss(y)
            except Exception:
                y_harm = y
        else:
            y_harm = y

        # 3. Chromagram
        chroma = librosa.feature.chroma_cqt(y=y_harm, sr=sr, hop_length=hop_length)
        if chroma.shape[1] == 0:
            return [{"error": "Could not compute chroma (audio too short or silent)."}]

        chroma_norm = chroma / (np.linalg.norm(chroma, axis=0, keepdims=True) + 1e-8)

        # 4. Template matching
        template_names = list(CHORD_TEMPLATES.keys())
        template_matrix = np.stack(
            [CHORD_TEMPLATES[name] for name in template_names],
            axis=1
        )  # (12, n_chords)

        sims = chroma_norm.T @ template_matrix  # (n_frames, n_chords)
        best_idx = np.argmax(sims, axis=1)
        best_scores = sims[np.arange(sims.shape[0]), best_idx]

        chord_labels = []
        for score, idx in zip(best_scores, best_idx):
            if score < chord_threshold:
                chord_labels.append('N')
            else:
                chord_labels.append(template_names[idx])

        # 5. Frames -> times
        times = librosa.frames_to_time(
            np.arange(len(chord_labels)),
            sr=sr,
            hop_length=hop_length
        )

        # 6. Group consecutive chords
        chords = []
        if len(chord_labels) > 0:
            last_chord = chord_labels[0]
            start_t = float(times[0])

            for i in range(1, len(chord_labels)):
                if chord_labels[i] != last_chord:
                    if last_chord != 'N':
                        chords.append({
                            "start_time": float(start_t),
                            "end_time": float(times[i]),
                            "chord_name": last_chord
                        })
                    last_chord = chord_labels[i]
                    start_t = float(times[i])

            if last_chord != 'N':
                chords.append({
                    "start_time": float(start_t),
                    "end_time": float(times[-1]),
                    "chord_name": last_chord
                })

        # 7. Smooth: enforce minimum duration
        chords = _enforce_min_duration(chords, min_chord_duration)

        return chords

    except Exception as e:
        # Show the real error in your console/logs
        import traceback
        traceback.print_exc()
        return [{"error": f"Chord detection error: {str(e)}"}]
    
def analyze_meta(file_path):
    """Extracts high-level metadata: BPM, Key, Loudness, etc."""
    # Load a 60-second mono segment for efficient analysis
    y, sr = librosa.load(file_path, duration=60) 

    # 1. Get BPM (Tempo)
    tempo, _ = librosa.beat.beat_track(y=y, sr=sr)
    
    # 2. Get Key and Mode (e.g., C# minor)
    # This is an advanced template-matching method for key detection.
    chroma = librosa.feature.chroma_stft(y=y, sr=sr)
    song_chroma_profile = np.sum(chroma, axis=1)

    # Define standard templates for major and minor keys (Krumhansl-Schmuckler profiles)
    major_template = np.array([6.35, 2.23, 3.48, 2.33, 4.38, 4.09, 2.52, 5.19, 2.39, 3.66, 2.29, 2.88])
    minor_template = np.array([6.33, 2.68, 3.52, 5.38, 2.60, 3.53, 2.54, 4.75, 3.98, 2.69, 3.34, 3.17])

    notes = ['C', 'C#', 'D', 'D#', 'E', 'F', 'F#', 'G', 'G#', 'A', 'A#', 'B']
    best_score = -1
    estimated_key = "N/A"

    # Correlate the song's profile with all 24 possible key templates
    for i in range(12):
        # Test against major template for the current root note
        rotated_major_template = np.roll(major_template, i)
        major_score = np.corrcoef(song_chroma_profile, rotated_major_template)[0, 1]
        if major_score > best_score:
            best_score = major_score
            estimated_key = f"{notes[i]} major"

        # Test against minor template for the current root note
        rotated_minor_template = np.roll(minor_template, i)
        minor_score = np.corrcoef(song_chroma_profile, rotated_minor_template)[0, 1]
        if minor_score > best_score:
            best_score = minor_score
            estimated_key = f"{notes[i]} minor"

    # 3. Get average loudness (RMS)
    # This was missing from the previous version, causing an error.
    rms = librosa.feature.rms(y=y)
    avg_rms = float(np.mean(rms))
    
    # 4. Get average spectral centroid (brightness)
    spectral_centroid = librosa.feature.spectral_centroid(y=y, sr=sr)
    avg_spectral_centroid = float(np.mean(spectral_centroid))

    # 5. Extract metadata tags (Artist, Title) using mutagen
    artist = None
    title = None
    try:
        audio_tags = mutagen.File(file_path, easy=True)
        if audio_tags:
            artist = audio_tags.get('artist', [None])[0]
            title = audio_tags.get('title', [None])[0]
    except Exception as e:
        print(f"--- [WARN] Could not read metadata tags with mutagen: {e}")

    # Fallback to filename parsing if tags are missing
    if not artist or not title:
        print("--- [INFO] Trying to parse artist/title from filename. ---")
        filename = os.path.basename(file_path)
        # A common format is "Artist - Title.mp3"
        if ' - ' in filename:
            try:
                parts = filename.rsplit('.', 1)[0].split(' - ', 1)
                if not artist: artist = parts[0].strip()
                if not title: title = parts[1].strip()
            except Exception: pass

    return {
        "bpm": round(float(tempo), 2), # Force float conversion
        "duration_seconds": float(librosa.get_duration(path=file_path)), # Get duration of the full file
        "estimated_key": estimated_key,
        "loudness_rms": round(avg_rms, 4),
        "brightness_spectral_centroid": round(avg_spectral_centroid, 2),
        "artist": artist,
        "title": title
    }

def analyze_lyrics(file_path, artist=None, title=None):
    """
    Analyzes lyrics by transcribing audio directly using the Gemini API.
    Returns: A dictionary with 'lyrics_lines' and 'chords': None.
    """
    gemini_api_key = os.environ.get("GEMINI_API_KEY")
    if not gemini_api_key:
        print("--- [WARN] GEMINI_API_KEY not set. Cannot perform lyric analysis. ---")
        return {"lyrics_lines": [], "chords": None}

    try:
        genai.configure(api_key=gemini_api_key)
        
        print(f"--- [INFO] Uploading {file_path} to Gemini for lyric analysis... ---")
        audio_file = genai.upload_file(path=file_path)
        
        # Wait for the file to be processed. This is a good practice for large files.
        while audio_file.state.name == "PROCESSING":
            print("... waiting for Gemini file processing...")
            time.sleep(10)
        
        if audio_file.state.name == "FAILED":
            print(f"--- [ERROR] Gemini file processing failed for {file_path}. ---")
            return {"lyrics_lines": [], "chords": None}
        
        print("--- [INFO] File uploaded. Transcribing with Gemini 1.5 Flash... ---")
        model = genai.GenerativeModel('gemini-2.5-flash')

        prompt = (
            "You are a highly specialized AI for music transcription, tasked with creating perfectly synchronized, time-aligned lyric data for a music application. Your output is machine-read, so precision and strict adherence to the format are paramount.\n\n"
            "**Your Task:**\n"
            "Listen to the provided audio file and generate a precise, time-aligned transcription of the lyrics for the ENTIRE song.\n\n"
            "**Output Requirements (Strictly Enforced):**\n"
            "1.  **Format:** Your entire response MUST be a single, valid JSON array of objects. Do NOT include any surrounding text, explanations, or markdown code fences (like ```json). Just the raw JSON array.\n"
            "2.  **Object Structure:** Each object in the array represents a single lyric line or instrumental segment and MUST contain three keys:\n"
            "    - `start`: The start time of the segment in seconds (float).\n"
            "    - `end`: The end time of the segment in seconds (float).\n"
            "    - `text`: The transcribed lyric text for the segment (string).\n"
            "3.  **Synchronization:** Timestamps (`start` and `end`) must be perfectly synchronized with the vocals in the audio. This is the most critical requirement.\n"
            "4.  **Contiguous Timeline:** The timeline must be complete and have no gaps. For any section without vocals (intros, outros, instrumental solos, long pauses), you MUST generate a segment with the text set to `'[Instrumental]'`.\n"
            "5.  **Language:** If the lyrics are not in English, provide a romanized (Roman English alphabet) transcription.\n"
            "6.  **Unintelligible Vocals:** If a word or phrase is completely unintelligible, use `[unintelligible]` in the text.\n\n"
            "**Example of a valid output snippet:**\n"
            "[\n"
            "  { \"start\": 0.0, \"end\": 10.5, \"text\": \"[Instrumental]\" },\n"
            "  { \"start\": 10.5, \"end\": 14.2, \"text\": \"Hello, is there anybody in there?\" },\n"
            "  { \"start\": 14.2, \"end\": 18.0, \"text\": \"Just nod if you can hear me.\" },\n"
            "  { \"start\": 18.0, \"end\": 19.5, \"text\": \"Is there anyone at home?\" }\n"
            "]"
        )

        response = model.generate_content([prompt, audio_file])
        
        # Clean the response to extract only the JSON part
        response_text = response.text.strip()
        json_start = response_text.find('[')
        json_end = response_text.rfind(']') + 1
        
        if json_start == -1 or json_end == 0:
            print(f"--- [ERROR] Gemini did not return a valid JSON array. Response: {response_text}")
            return {"lyrics_lines": [], "chords": None}
            
        json_string = response_text[json_start:json_end]
        
        lyrics_lines = json.loads(json_string)
        
        # Basic validation of the returned structure
        if not isinstance(lyrics_lines, list) or (lyrics_lines and not isinstance(lyrics_lines[0], dict)):
             print(f"--- [ERROR] Gemini returned JSON but in the wrong format. Response: {json_string}")
             return {"lyrics_lines": [], "chords": None}

        print(f"--- [INFO] Gemini transcription successful. Found {len(lyrics_lines)} lines. ---")
        
        # The caller expects chords to be None to trigger fallback chord analysis.
        return {"lyrics_lines": lyrics_lines, "chords": None}

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"--- [ERROR] An error occurred during Gemini lyric analysis: {e}")
        return {"lyrics_lines": [], "chords": None}
        
def merge_lyrics_and_chords(lyrics_data, chords_data):
    """
    Merges timestamped lyrics with timestamped chords.
    For each line of lyrics, it finds the chord playing at the beginning of the line.
    """
    if not lyrics_data or not chords_data:
        return []

    merged_data = []
    chord_idx = 0

    for line in lyrics_data:
        # Find the chord that is active at the start of the lyric line
        current_chord = "N" # No Chord
        
        # Advance chord_idx to find the relevant chord
        while chord_idx < len(chords_data) and chords_data[chord_idx]['end_time'] < line['start']:
            chord_idx += 1
        
        if chord_idx < len(chords_data) and chords_data[chord_idx]['start_time'] <= line['start'] < chords_data[chord_idx]['end_time']:
            current_chord = chords_data[chord_idx]['chord_name']
        
        merged_data.append({
            "start": line['start'],
            "end": line['end'],
            "text": line['text'],
            "chord": current_chord
        })

    return merged_data

def create_lyrics_doc(song_title, merged_data, output_doc_path):
    """
    Creates a .docx file with chords placed above the lyrics.
    """
    try:
        doc = Document()
        doc.add_heading(song_title, level=1)
        doc.add_paragraph() # Add some space

        p = doc.add_paragraph()
        for item in merged_data:
            chord_run = p.add_run(f"{item['chord']}\n")
            chord_run.bold = True
            chord_run.font.name = 'Courier New'
            lyric_run = p.add_run(f"{item['text']}\n\n")
            lyric_run.font.name = 'Arial'

        doc.save(output_doc_path)
        return output_doc_path
    except Exception as e:
        print(f"--- [ERROR] Failed to create .docx file: {e}")
        return None



    
    
def analyze_notes_for_stems(stems_dict):
    """
    Runs note detection on all relevant stems (excluding drums) and returns a
    dictionary of the results.
    """
    all_notes = {}
    
    # Define which stems are melodic/harmonic and should be analyzed for notes
    stems_to_process = ['vocals', 'bass', 'piano', 'guitar', 'other']
    
    for stem_name in stems_to_process:
        if stem_name in stems_dict and os.path.exists(stems_dict[stem_name]):
            print(f"--- [Internal] Analyzing notes for '{stem_name}' stem ---")
            # This reuses the existing single-file analysis function
            notes = analyze_notes_basic_pitch(stems_dict[stem_name])
            all_notes[stem_name] = notes
        else:
            print(f"--- [WARN] Stem '{stem_name}' not found or file is missing. Skipping note analysis. ---")
            all_notes[stem_name] = [] # Return empty list for missing stems
            
    return all_notes

def analyze_notes_basic_pitch(file_path):
    """
    Uses Spotify's Basic Pitch to detect MIDI notes.
    """
    try:
        # predict returns: model_output, midi_data, note_events
        _, _, note_events = predict(file_path)
        
        notes_json = []
        for note in note_events:
            # CRITICAL FIX: Cast numpy types to Python native types
            notes_json.append({
                "start": float(note[0]),
                "end": float(note[1]),
                "pitch": int(note[2]),    # Cast numpy.int64 to int
                "velocity": float(note[3]) # Cast numpy.float32 to float
            })
        return notes_json
    except Exception as e:
        print(f"Error in Basic Pitch: {e}")
        return []

def separate_stems(file_path, output_dir):
    """
    Calls Demucs to perform a 6-stem separation using the htdemucs_6s model.
    This model separates audio into: vocals, bass, drums, piano, guitar, and other.
    """
    filename_base = os.path.basename(file_path).split('.')[0]
    model_6s = "htdemucs_6s"
    print(f"Starting Demucs 6-stem separation ({model_6s}) for {file_path}...")
    command_6s = ["demucs", "-n", model_6s, "--out", output_dir, file_path]
    subprocess.run(command_6s, check=True)

    print("Stem separation complete.")

    # Define path to the output directory of the model
    path_6s = os.path.join(output_dir, model_6s, filename_base)
    
    # Create a dictionary pointing to each separated stem file.
    return {
        "vocals": os.path.join(path_6s, "vocals.wav"),
        "bass": os.path.join(path_6s, "bass.wav"),
        "drums": os.path.join(path_6s, "drums.wav"),
        "piano": os.path.join(path_6s, "piano.wav"),
        "guitar": os.path.join(path_6s, "guitar.wav"),
        "other": os.path.join(path_6s, "other.wav"),
    }